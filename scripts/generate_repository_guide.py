#!/usr/bin/env python3
"""Generate the repository's comprehensive Microsoft Word reading guide."""

from __future__ import annotations

import re
import subprocess
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Sales-App-SDLC-Repository-Guide.docx"
EXCLUDED = {OUTPUT.relative_to(ROOT).as_posix()}

SERVICE_SUMMARIES = {
    "customer-service": (
        "Java 21 / Spring Boot",
        "PostgreSQL through Spring Data JPA",
        "Owns customer identity and lifecycle. Its controller saves the customer and an outbox event "
        "inside one transaction, demonstrating the transactional-outbox boundary.",
    ),
    "product-catalog-service": (
        "TypeScript / NestJS",
        "In-memory Map in the skeleton; target is Cosmos DB + Azure AI Search",
        "Creates draft products, publishes or updates them, and records product-domain events.",
    ),
    "pricing-service": (
        "Python / FastAPI",
        "In-memory dictionaries in the skeleton; target is PostgreSQL + Redis",
        "Maintains prices and promotions and calculates quote subtotal, discount, fixed 8% tax, and total.",
    ),
    "inventory-service": (
        "Go / Gin",
        "Mutex-protected in-memory store; target is Cosmos DB",
        "Tracks on-hand and reserved stock, creates reservations atomically, and releases them for compensation.",
    ),
    "order-service": (
        "C# / ASP.NET Core minimal APIs",
        "In-memory dictionary; target is Azure SQL",
        "Acts as the quote-to-cash coordinator. It creates orders, records reservation/payment IDs, confirms, "
        "and marks compensated orders.",
    ),
    "payment-service": (
        "Kotlin / Ktor",
        "In-memory payment ledger; target is Azure SQL",
        "Authorizes idempotently, simulates provider failure, captures, refunds, and emits payment events.",
    ),
    "notification-service": (
        "Rust / Axum",
        "Arc<Mutex<...>> in-memory store; target is Cosmos DB",
        "Accepts notification requests and records requested/delivered events before exposing status by ID.",
    ),
    "sales-analytics-service": (
        "Scala / Play Framework",
        "In-memory ListBuffer; target is Azure Data Explorer / Fabric",
        "Serves KPI snapshots and stores forecasts while producing a forecast-updated event.",
    ),
}

CORE_READING_ORDER = [
    ("README.md", "Start with intent, boundaries, the system picture, and delivery model."),
    ("docs/architecture.md", "Learn the target Azure architecture and strict tool ownership boundaries."),
    ("contracts/openapi/*.yaml", "Read the synchronous API promises before implementations."),
    ("contracts/events/*.json", "Read the asynchronous CloudEvent envelope and payload definitions."),
    ("apps/web/app/page.js", "See the intentionally thin user-facing entry point."),
    ("services/customer-service", "Follow a persisted write plus transactional outbox."),
    ("services/product-catalog-service", "See NestJS controller-to-provider dependency injection."),
    ("services/pricing-service/app/main.py", "Read the smallest complete quote-calculation flow."),
    ("services/inventory-service", "Follow concurrency-safe reserve/release behavior."),
    ("services/payment-service", "Follow idempotent authorization and payment state changes."),
    ("services/order-service", "Connect inventory and payment IDs through the saga coordinator."),
    ("services/notification-service", "See a downstream event-oriented delivery boundary."),
    ("services/sales-analytics-service", "See the read/insight side of the platform."),
    ("docker-compose.yml", "Understand local process wiring and ports."),
    ("charts/sales-service", "Read the reusable Kubernetes workload contract."),
    ("gitops/argocd then gitops/apps", "Trace desired state from ApplicationSet to environment values."),
    ("infrastructure/environments then infrastructure/modules", "Trace Azure composition into resources."),
    (".github/workflows", "Finish with verification, packaging, signing, and release automation."),
]

FLOW_STEPS = [
    ("1. Enter", "Browser/mobile traffic reaches Front Door + WAF, API Management, then the AKS Gateway. "
     "Those managed edge resources are architecture targets; the local skeleton starts at each service route."),
    ("2. Identify customer", "Customer Service creates or reads the buyer. The create transaction writes both "
     "the customer row and CustomerCreated outbox row."),
    ("3. Select product", "Product Catalog returns a published product/SKU and records publish/change events."),
    ("4. Price", "Pricing resolves SKU prices, validates currency and promotions, then computes subtotal, "
     "discount, tax, and total."),
    ("5. Reserve", "Inventory locks its store, checks available = onHand - reserved, increments reserved stock, "
     "and returns a reservation ID."),
    ("6. Create/confirm order", "Order Service records customer, lines, and idempotency key. Its saga associates "
     "the inventory reservation and payment authorization before confirming."),
    ("7. Pay", "Payment Service returns an existing authorization for a repeated idempotency key; otherwise it "
     "records an authorized or failed payment and provider reference."),
    ("8. Compensate on failure", "The intended saga releases inventory and reverses payment if a later step "
     "fails. The skeleton exposes release/refund and order compensation primitives, but does not yet orchestrate "
     "remote calls or a Service Bus consumer."),
    ("9. Notify and analyze", "Notification records requested/delivered events. Analytics exposes KPI and "
     "forecast endpoints. Production intent is Service Bus delivery, not direct shared-memory access."),
    ("10. Observe", "Workloads export OTLP telemetry to the collector, which forwards traces/metrics/logs to the "
     "configured backends in the target platform."),
]

PATH_PURPOSES = {
    ".github/": "Repository ownership, dependency automation, pull-request verification, releases, and Terraform automation.",
    "ansible/": "Configuration of non-Kubernetes hosts such as self-hosted runners and utility VMs.",
    "apps/web/": "Minimal Next.js user experience and container definition.",
    "charts/sales-service/": "Golden reusable Helm chart encoding workload security, availability, networking, and observability defaults.",
    "contracts/events/": "Versioned asynchronous CloudEvent envelope and payload schemas.",
    "contracts/openapi/": "Versioned synchronous HTTP contracts for all eight services.",
    "docs/": "Cross-cutting architecture, reliability, security, testing, and operating guidance.",
    "gitops/": "Argo CD bootstrap and environment-specific desired-state values.",
    "infrastructure/environments/": "Terraform composition roots for dev, test, stage, and prod.",
    "infrastructure/modules/": "Reusable Azure resource modules for networking, AKS, ACR, Key Vault, data, messaging, and observability.",
    "policies/": "Kubernetes admission policy enforcing baseline workload security.",
    "scripts/": "Repository-level validation, test dispatch, load testing, and documentation helpers.",
    "services/": "Eight independently buildable bounded-context walking skeletons.",
}

FILE_NAME_PURPOSES = {
    "Dockerfile": "Multi-stage or service-specific container build instructions.",
    "README.md": "Human-facing overview and operating instructions for this scope.",
    "Chart.yaml": "Helm package metadata and dependency declaration.",
    "values.yaml": "Default Helm values or service-specific chart overrides.",
    "openapi.yaml": "OpenAPI contract describing supported HTTP operations and schemas.",
    "runbook.md": "Operational triage, health, alert, and rollback guidance.",
    "threat-model.md": "Service assets, trust boundaries, threats, and mitigations.",
    "contract-test.sh": "Dispatches contract validation for this service.",
    "lint.sh": "Runs language-appropriate static analysis.",
    "test.sh": "Runs language-appropriate service tests.",
    "sonar-project.properties": "SonarQube project identity and source/test analysis settings.",
}


def tracked_files() -> list[Path]:
    raw = subprocess.check_output(["git", "ls-files", "-z"], cwd=ROOT)
    return [
        ROOT / item.decode()
        for item in raw.split(b"\0")
        if item and item.decode() not in EXCLUDED
    ]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=40, start=70, bottom=40, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update field in Word"
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in (
        ("Title", 28, "17365D"),
        ("Heading 1", 19, "17365D"),
        ("Heading 2", 14, "2F5597"),
        ("Heading 3", 11, "4472C4"),
        ("Heading 4", 9, "5B9BD5"),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    if "Code" not in styles:
        code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code"]
    code.font.name = "Cascadia Mono"
    code.font.size = Pt(7)
    code.font.color.rgb = RGBColor(31, 31, 31)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.text = "Sales-App-SDLC  |  Repository Reading Guide"
    header.style = styles["Caption"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    add_field(footer, "PAGE")
    footer.add_run(" of ")
    add_field(footer, "NUMPAGES")


def add_title(doc: Document, files: list[Path]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("Sales-App-SDLC")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(23, 54, 93)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Repository Reading Guide")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(47, 85, 151)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Architecture, end-to-end flow, recommended reading order,\n"
              "and an exhaustive line-by-line companion").italic = True

    text_lines = sum(len(path.read_text(encoding="utf-8", errors="replace").splitlines()) for path in files)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    p.add_run(f"Repository snapshot: {len(files)} tracked text files · {text_lines:,} physical lines\n")
    p.add_run(f"Generated {date.today().isoformat()}").italic = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(35)
    note.add_run(
        "Audience: engineers, reviewers, operators, security staff, and technical stakeholders "
        "who need to navigate this polyglot walking skeleton."
    )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-4" \\h \\z \\u')
    doc.add_paragraph(
        "If Word does not refresh this table automatically, right-click it and choose Update Field. "
        "The headings remain navigable in Word's Navigation pane."
    )
    doc.add_page_break()


def add_scope(doc: Document, files: list[Path]) -> None:
    doc.add_heading("1. How to use this guide", level=1)
    doc.add_paragraph(
        "This is a monorepo walking skeleton: it proves boundaries, contracts, delivery mechanics, and one "
        "representative path through eight services. It is not a finished production commerce platform. "
        "Many service stores and event lists are deliberately in memory, while the documents describe their "
        "target managed Azure data services and Service Bus integration."
    )
    doc.add_heading("What “line by line” means here", level=2)
    doc.add_paragraph(
        "Part I explains intent and flow at human scale. Part II inventories every tracked file. Part III then "
        "lists every physical line of every tracked text file with a concise interpretation. Blank lines are "
        "included because they reveal block boundaries. Dependency lockfiles are also included, but their "
        "machine-generated records are identified as resolution data rather than hand-authored application logic."
    )
    doc.add_paragraph(
        "Do not read Part III from page one to the end. Use the recommended path below, search for a file name, "
        "then use the adjacent line numbers to relate an explanation to the repository snapshot."
    )

    suffixes = Counter(path.suffix or "(no extension)" for path in files)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Snapshot measure"
    table.rows[0].cells[1].text = "Value"
    rows = [
        ("Tracked text files", str(len(files))),
        ("Physical lines", f"{sum(len(p.read_text(errors='replace').splitlines()) for p in files):,}"),
        ("Largest file", max(files, key=lambda p: len(p.read_text(errors="replace").splitlines())).relative_to(ROOT).as_posix()),
        ("Most common extensions", ", ".join(f"{k}: {v}" for k, v in suffixes.most_common(8))),
    ]
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = key, value


def add_system_map(doc: Document) -> None:
    doc.add_heading("2. System map and responsibility boundaries", level=1)
    doc.add_paragraph(
        "The central design rule is ownership: each bounded context owns its code and database; contracts are "
        "versioned; synchronous REST is used for the walking skeleton; asynchronous integration targets Azure "
        "Service Bus; and no service reaches into another service's database."
    )
    doc.add_heading("Runtime path", level=2)
    runtime = (
        "Browser / mobile client\n"
        "  → Azure Front Door Premium + WAF\n"
        "  → Azure API Management\n"
        "  → AKS Gateway / HTTPRoute\n"
        "  → domain service REST endpoint\n"
        "  → service-owned state + outbox\n"
        "  → Azure Service Bus topic/queue\n"
        "  → Notification / Analytics consumers\n"
        "  ↘ OpenTelemetry Collector → Azure Monitor / Prometheus / Grafana"
    )
    p = doc.add_paragraph(style="Code")
    p.add_run(runtime)

    doc.add_heading("Delivery path", level=2)
    delivery = (
        "Source change\n"
        "  → pull-request path filter\n"
        "  → language lint/test + contract validation + Sonar/CodeQL/Gitleaks\n"
        "  → non-release container build + Trivy + Syft + Helm lint\n"
        "  → merge to main\n"
        "  → BuildKit image + SBOM/provenance + Cosign signature\n"
        "  → immutable ACR digest + signed OCI Helm chart\n"
        "  → GitOps values PR\n"
        "  → Argo CD reconciliation\n"
        "  → same digest promoted dev → test → stage → prod"
    )
    p = doc.add_paragraph(style="Code")
    p.add_run(delivery)

    doc.add_heading("Infrastructure ownership", level=2)
    boundaries = [
        ("Terraform", "Creates Azure platform resources. It must not continually mutate application desired state."),
        ("Helm", "Packages a service into Kubernetes resources with common secure defaults."),
        ("Argo CD", "Reconciles Git-declared Helm values into AKS; rollback is a Git revert."),
        ("GitHub Actions", "Verifies, builds, scans, signs, publishes, and proposes GitOps changes."),
        ("Ansible", "Configures runners and utility VMs only; it is not an AKS deployment tool."),
        ("Kyverno", "Rejects Kubernetes workloads that violate baseline security policy."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Tool"
    table.rows[0].cells[1].text = "Owned responsibility"
    for tool, ownership in boundaries:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = tool, ownership


def add_business_flow(doc: Document) -> None:
    doc.add_heading("3. End-to-end business flow", level=1)
    doc.add_paragraph(
        "The intended walking-skeleton journey is customer → product → quote → inventory → order → payment → "
        "notification → analytics. The current code exposes the primitives but does not yet contain a single "
        "cross-service orchestrator or working Service Bus publisher. That distinction matters when reading it."
    )
    for title, detail in FLOW_STEPS:
        doc.add_heading(title, level=2)
        doc.add_paragraph(detail)

    doc.add_heading("Happy-path object identifiers", level=2)
    doc.add_paragraph(
        "Customer Service creates customerId. Product Catalog supplies SKU. Pricing returns money totals. "
        "Inventory returns reservationId. Order Service creates orderId and later stores reservationId and "
        "paymentId. Payment returns paymentId and providerReference. Notification returns notificationId. "
        "These IDs, not shared database rows, connect the domains."
    )

    doc.add_heading("Important implementation gaps to recognize", level=2)
    gaps = [
        "Most stores are process-local memory and lose data on restart; Customer Service is the main persistence example.",
        "Outbox/event collections are not drained by a Service Bus publisher in this snapshot.",
        "OrderSaga records supplied reservation/payment IDs but does not call Inventory or Payment itself.",
        "Health endpoints generally report UP without checking dependencies.",
        "Authentication/authorization, production validation depth, migrations, retries, dead-letter handling, and "
        "distributed tracing propagation are architecture intent rather than complete implementation.",
        "The fixed pricing tax rate and immediate notification delivery are demonstration behavior, not production rules.",
    ]
    for gap in gaps:
        doc.add_paragraph(gap, style="List Bullet")


def add_service_tour(doc: Document) -> None:
    doc.add_heading("4. Service-by-service tour", level=1)
    for name, (stack, state, purpose) in SERVICE_SUMMARIES.items():
        doc.add_heading(name, level=2)
        doc.add_paragraph(f"Stack: {stack}. State: {state}.")
        doc.add_paragraph(purpose)
        details = {
            "customer-service": (
                "Read CustomerServiceApplication → Customer entity/repository → CustomerController → OutboxEvent. "
                "The @Transactional create/update methods are the key consistency boundary."
            ),
            "product-catalog-service": (
                "Read main.ts → AppModule → ProductsController → ProductsService → product.ts. NestJS decorators "
                "wire routes and dependency injection; the service owns state transitions."
            ),
            "pricing-service": (
                "Read app/main.py top to bottom. Pydantic models define the contract, dictionaries model state, "
                "decorators define routes, and calculate_quote holds the central algorithm."
            ),
            "inventory-service": (
                "Read cmd/server/main.go → internal/api/http.go → internal/domain/store.go. HTTP concerns stay "
                "outside the mutex-protected stock/reservation rules."
            ),
            "order-service": (
                "Read Program.cs → Api/OrderEndpoints.cs → Domain/Orders.cs → Saga/OrderSaga.cs. The endpoint layer "
                "creates commands; OrderStore records state/events; OrderSaga changes lifecycle state."
            ),
            "payment-service": (
                "Read Application.kt → domain/Payment.kt. Ktor maps HTTP to PaymentLedger; the ledger enforces "
                "idempotent authorization and state transitions."
            ),
            "notification-service": (
                "Read src/main.rs in blocks: data types → handlers → router → server → test. Arc<Mutex<Store>> "
                "makes shared mutable state safe enough for this single-process demonstration."
            ),
            "sales-analytics-service": (
                "Read conf/routes → AnalyticsController → models/Analytics.scala → AnalyticsService. Play routes "
                "enter the controller, JSON validation creates a Forecast, and the service stores it and emits an event."
            ),
        }[name]
        doc.add_paragraph(details)


def add_reading_order(doc: Document) -> None:
    doc.add_heading("5. Recommended way to read the repository", level=1)
    doc.add_paragraph(
        "Read contracts and boundaries before implementation details. For every service, use the same loop: README "
        "→ OpenAPI → entry point → HTTP/controller layer → domain/service layer → tests → Dockerfile → Helm values "
        "→ runbook/threat model. This makes the eight languages comparable."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    headers = ("Step", "Path", "Question to answer")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for index, (path, reason) in enumerate(CORE_READING_ORDER, 1):
        cells = table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = path
        cells[2].text = reason

    doc.add_heading("A repeatable code-reading checklist", level=2)
    checks = [
        "Where does execution enter (main function, framework bootstrap, route table)?",
        "Which route and schema accept the request?",
        "Where are domain validation and state transitions implemented?",
        "Where is state stored, and is it durable or only in memory?",
        "Which event is recorded, and is it transactionally coupled to state?",
        "How do failure, idempotency, retry, and compensation behave?",
        "Which tests prove the behavior?",
        "How is the process built, configured, observed, deployed, and operated?",
    ]
    for item in checks:
        doc.add_paragraph(item, style="List Number")


def add_local_and_delivery(doc: Document) -> None:
    doc.add_heading("6. Running, testing, and shipping", level=1)
    doc.add_heading("Local stack", level=2)
    doc.add_paragraph(
        "Copy .env.example to .env, validate contracts, run language tests, then use docker compose up --build. "
        "The compose file maps web to 3000 and services to host ports 8081–8088 while containers listen on 8080. "
        "PostgreSQL, Redis, MongoDB, and the OTLP collector provide representative local dependencies."
    )
    p = doc.add_paragraph(style="Code")
    p.add_run(
        "make contracts\n"
        "make test-python\n"
        "make test-go\n"
        "make test-rust\n"
        "cp .env.example .env\n"
        "docker compose up --build"
    )
    doc.add_heading("Pull request", level=2)
    doc.add_paragraph(
        "pull-request.yml detects changed paths and invokes reusable-service-ci.yml for affected services. "
        "The reusable workflow dispatches language-specific scripts, validates contracts, performs analysis and "
        "supply-chain checks, builds a non-release image, and lints the Helm package."
    )
    doc.add_heading("Release and GitOps", level=2)
    doc.add_paragraph(
        "release.yml invokes reusable-service-release.yml. The release path builds once, scans and signs immutable "
        "artifacts, and updates GitOps by digest. Argo CD's app-of-apps/ApplicationSet owns reconciliation. "
        "Environment files differ primarily in replica count, resource sizing, autoscaling, and immutable image digest."
    )
    doc.add_heading("Terraform", level=2)
    doc.add_paragraph(
        "Each infrastructure/environments/<env>/main.tf composes the same modules. backend.tf selects remote state; "
        "variables.tf declares inputs; outputs.tf exposes identifiers. Modules should be read from network and "
        "identity outward: network → ACR/Key Vault/observability/data/messaging → AKS."
    )


def file_purpose(rel: str) -> str:
    path = Path(rel)
    if path.name in FILE_NAME_PURPOSES:
        base = FILE_NAME_PURPOSES[path.name]
    else:
        base = ""
    if rel == "README.md":
        return "Top-level system intent, service catalogue, runtime flow, repository map, and delivery overview."
    if rel == "docker-compose.yml":
        return "Local multi-container topology, dependencies, environment variables, and host port mappings."
    if rel == "Makefile":
        return "Convenience entry points for contract checks and selected language test suites."
    if rel.startswith("services/"):
        parts = path.parts
        service = parts[1] if len(parts) > 1 else ""
        service_text = SERVICE_SUMMARIES.get(service, ("", "", "Service implementation."))[2]
        return f"{base + ' ' if base else ''}{service_text}".strip()
    for prefix, purpose in PATH_PURPOSES.items():
        if rel.startswith(prefix):
            return f"{base + ' ' if base else ''}{purpose}".strip()
    if rel.endswith(".lock") or rel.endswith("package-lock.json") or rel.endswith("go.sum"):
        return "Machine-generated dependency resolution lock/checksum data; review changes, do not read as application flow."
    return base or "Repository configuration or documentation supporting the platform."


def add_inventory(doc: Document, files: list[Path]) -> None:
    doc.add_page_break()
    doc.add_heading("Part II — Complete file inventory", level=1)
    doc.add_paragraph(
        "Files are grouped by top-level directory. Every file listed here has a matching line-by-line section in Part III."
    )
    groups: dict[str, list[Path]] = {}
    for path in files:
        rel = path.relative_to(ROOT)
        groups.setdefault(rel.parts[0], []).append(path)
    for group, paths in sorted(groups.items()):
        doc.add_heading(group, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "File"
        table.rows[0].cells[1].text = "Lines"
        table.rows[0].cells[2].text = "Responsibility"
        for path in paths:
            rel = path.relative_to(ROOT).as_posix()
            lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
            cells = table.add_row().cells
            cells[0].text = rel
            cells[1].text = str(len(lines))
            cells[2].text = file_purpose(rel)


def yaml_explanation(stripped: str, indent: int) -> str:
    if stripped.startswith("- "):
        return f"Adds a list item at YAML indentation level {indent // 2}."
    match = re.match(r"([^:#]+):(?:\s*(.*))?$", stripped)
    if match:
        key, value = match.group(1).strip(), (match.group(2) or "").strip()
        if value:
            return f"Sets YAML key “{key}” to {value}; indentation scopes it under its parent."
        return f"Starts YAML mapping “{key}”; the indented lines below belong to it."
    if "{{" in stripped:
        return "Evaluates Helm/Actions template expressions to produce environment-specific configuration."
    return "Continues the current YAML value or template block."


def explain_line(rel: str, number: int, line: str, previous: str) -> str:
    stripped = line.strip()
    suffix = Path(rel).suffix.lower()
    name = Path(rel).name

    if not stripped:
        return "Blank separator: visually ends one logical block before the next."
    if rel.endswith(("package-lock.json", "Cargo.lock", "go.sum")):
        if stripped.startswith(("{", "}", "[[", "name =", "version =", '"node_modules/', '"version"', '"integrity"')):
            return "Machine-generated dependency resolution metadata; it pins an exact package or checksum."
        return "Continues machine-generated dependency lock/checksum data."
    if stripped.startswith(("#", "//", "/*", "*", "<!--")):
        return "Comment or documentation text for readers; it does not directly execute."
    if suffix == ".md":
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            return f"Markdown heading level {level}; starts a new documentation topic."
        if stripped.startswith("```"):
            return "Opens or closes a fenced example/code block."
        if stripped.startswith("|"):
            return "Defines a Markdown table row."
        if stripped.startswith(("-", "*")):
            return "Adds a documentation list item."
        if re.match(r"\d+\.", stripped):
            return "Adds a numbered process or checklist step."
        return "Explains the design, operation, or expected behavior in prose."
    if suffix in {".yaml", ".yml"}:
        return yaml_explanation(stripped, len(line) - len(line.lstrip()))
    if suffix == ".json":
        if stripped in {"{", "}", "[", "],", "}", "},"}:
            return "Opens or closes a JSON object/array structure."
        match = re.match(r'"([^"]+)":\s*(.*)', stripped.rstrip(","))
        if match:
            return f"Defines JSON property “{match.group(1)}” with its configured value or nested structure."
        return "Continues a JSON array/object value."
    if suffix == ".tf":
        if match := re.match(r'(resource|module|data|variable|output|locals|terraform|provider)\s+"?([^"{ ]*)"?', stripped):
            kind, identifier = match.groups()
            return f"Starts Terraform {kind} block “{identifier}”; nested arguments declare desired infrastructure."
        if stripped == "}":
            return "Closes the current Terraform block."
        if "=" in stripped:
            key = stripped.split("=", 1)[0].strip()
            return f"Assigns Terraform argument/local “{key}”; expressions are resolved during plan/apply."
        return "Continues the current Terraform expression or collection."
    if name == "Dockerfile":
        instruction = stripped.split(maxsplit=1)[0].upper()
        meanings = {
            "FROM": "Selects a container build stage/base image.",
            "WORKDIR": "Sets the working directory for following build instructions.",
            "COPY": "Copies repository or prior-stage files into the image.",
            "RUN": "Executes a build-time command and commits its filesystem result.",
            "ENV": "Sets an image/runtime environment default.",
            "EXPOSE": "Documents the port expected to receive traffic.",
            "USER": "Changes the account used by subsequent steps/runtime.",
            "ENTRYPOINT": "Defines the executable launched when the container starts.",
            "CMD": "Defines default runtime arguments/command.",
        }
        return meanings.get(instruction, "Continues the container image build definition.")
    if name == "Makefile" or suffix == ".mk":
        if not line.startswith(("\t", " ")) and ":" in stripped:
            return f"Defines Make target “{stripped.split(':', 1)[0]}” and its prerequisites."
        if line.startswith("\t"):
            return "Runs this shell recipe when the enclosing Make target is invoked."
        return "Sets or documents Make build automation."
    if suffix in {".sh"}:
        if stripped.startswith("#!"):
            return "Selects the shell interpreter."
        if stripped.startswith(("set ", "set -")):
            return "Enables strict shell error-handling behavior."
        if stripped.startswith(("if ", "elif ", "else", "fi")):
            return "Controls conditional shell execution."
        return "Executes or continues a shell command in this validation/test helper."
    if re.match(r"^(import|from|using|use|package)\b", stripped):
        return "Declares the namespace or imports a dependency used by this file."
    if stripped.startswith(("@", "#[")):
        return "Framework annotation/attribute: supplies routing, validation, persistence, DI, serialization, or test metadata."
    if re.search(r"\b(class|record|struct|enum|interface|object|type)\b", stripped):
        return "Declares a type that models domain data, behavior, configuration, or an API shape."
    if re.match(r"^(async\s+)?(def|fun|func|fn)\b", stripped) or re.search(
        r"\b(public|private|protected|static|async)\b.*\([^;]*\)\s*(?:\{|=>|=)?$", stripped
    ):
        return "Declares a callable unit; its parameters are inputs and the following block implements behavior."
    if stripped.startswith(("if ", "if(", "if (", "when ", "match ", "switch ")):
        return "Branches behavior according to a domain, validation, or error condition."
    if stripped.startswith(("for ", "for(", "while ", "forEach")):
        return "Iterates over a collection or repeats the enclosed behavior."
    if stripped.startswith(("return ", "return;", "Ok(", "Created(", "BadRequest(")):
        return "Returns the computed value or HTTP result to the caller."
    if any(token in stripped for token in (".get(", ".post(", ".put(", "MapGet(", "MapPost(", "@app.", "@Get", "@Post", "@Put")):
        return "Registers or invokes HTTP-facing behavior for the indicated route/operation."
    if "events" in stripped.lower() or "outbox" in stripped.lower():
        return "Reads, writes, or defines domain-event/outbox data for asynchronous integration."
    if stripped in {"}", "};", ")", ");", "]", "],", "};"}:
        return "Closes the current code block, call, object, or collection."
    if "=" in stripped:
        lhs = stripped.split("=", 1)[0].strip()
        return f"Assigns or initializes “{lhs[-45:]}” for use by the surrounding behavior."
    if stripped.startswith(("app.", "engine.", "router.", ".route(")):
        return "Configures framework middleware, routing, or application startup."
    if stripped.startswith(("throw ", "raise ", "error(", "require(", "assert")):
        return "Stops or verifies execution when an expected invariant is not satisfied."
    if suffix in {".csproj", ".sln", ".xml"}:
        return "Declares project/build metadata, dependency references, or XML structure."
    if suffix in {".properties", ".conf", ".cfg", ".ini"}:
        return "Sets a tool, framework, route, or runtime configuration value."
    if suffix in {".js", ".ts", ".tsx", ".py", ".go", ".rs", ".java", ".kt", ".kts", ".cs", ".scala"}:
        return "Continues the surrounding declaration, expression, function call, or data structure."
    if previous.strip().endswith("\\"):
        return "Continues the command from the preceding physical line."
    return "Defines or continues repository configuration for this file's responsibility."


def add_line_by_line(doc: Document, files: list[Path]) -> None:
    doc.add_page_break()
    doc.add_heading("Part III — Exhaustive line-by-line companion", level=1)
    doc.add_paragraph(
        "Each table is tied to this repository snapshot. “Code / content” is the exact physical line, including "
        "line numbering supplied by this guide (not stored in the file). Explanations describe the line's role in "
        "its local block. For semantics that span many lines, read the table together with Parts I and II."
    )

    for index, path in enumerate(files, 1):
        rel = path.relative_to(ROOT).as_posix()
        lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
        doc.add_heading(f"{index}. {rel}", level=2)
        doc.add_paragraph(file_purpose(rel))
        doc.add_paragraph(f"{len(lines)} physical lines · tracked text file", style="Caption")
        if not lines:
            doc.add_paragraph("This file is intentionally empty; its presence marks a package/module boundary.")
            continue

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(0.42)
        table.columns[1].width = Inches(3.65)
        table.columns[2].width = Inches(3.65)
        header = table.rows[0]
        set_repeat_table_header(header)
        for cell, text in zip(header.cells, ("Line", "Code / content", "What it does")):
            cell.text = text
            shade(cell, "D9EAF7")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(7)

        previous = ""
        for number, line in enumerate(lines, 1):
            cells = table.add_row().cells
            cells[0].text = str(number)
            code_p = cells[1].paragraphs[0]
            code_p.style = doc.styles["Code"]
            code_p.add_run(line if line else "⟨blank⟩")
            exp_p = cells[2].paragraphs[0]
            exp_p.add_run(explain_line(rel, number, line, previous))
            for cell in cells:
                set_cell_margin(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(6.5)
            previous = line

        if index < len(files):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A — Glossary", level=1)
    terms = [
        ("AKS", "Azure Kubernetes Service."),
        ("ACR", "Azure Container Registry."),
        ("Bounded context", "A business capability with its own model, API, data, and ownership boundary."),
        ("CloudEvent", "A standard event envelope carrying type, source, ID, time, and domain payload."),
        ("Compensation", "A semantic undo step, such as releasing inventory after payment fails."),
        ("Digest", "Immutable content hash identifying the exact container image promoted across environments."),
        ("GitOps", "Git is desired-state source; a controller reconciles the cluster to it."),
        ("HPA / KEDA", "CPU/resource-based and event-driven Kubernetes autoscaling mechanisms."),
        ("Idempotency key", "Client-supplied key that makes a repeated command return the original result."),
        ("Outbox", "Events committed with domain state and published later, avoiding a database/message split-brain."),
        ("PDB", "PodDisruptionBudget; limits voluntary simultaneous pod disruption."),
        ("Saga", "Multi-service business transaction made from local commits and compensating actions."),
        ("SBOM", "Software Bill of Materials describing artifact dependencies."),
        ("Workload Identity", "Federated pod-to-Azure authentication without embedded long-lived secrets."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Meaning in this repository"
    for term, meaning in terms:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = term, meaning

    doc.add_heading("Appendix B — Fast navigation by task", level=1)
    tasks = [
        ("Change an HTTP API", "Contract OpenAPI → controller/router → domain/service → unit test → service docs OpenAPI."),
        ("Change an event", "contracts/events schema → producer/outbox → consumer intent → compatibility validation."),
        ("Change Kubernetes runtime", "golden chart values/schema → template → per-service values → GitOps environment values."),
        ("Change Azure infrastructure", "environment root variables/main → module → outputs → Terraform PR workflow/policy."),
        ("Investigate production behavior", "service runbook → health/metrics → Helm probes/ServiceMonitor → SRE guidance."),
        ("Investigate CI/release", "pull-request/release workflow → reusable workflow → service scripts/Dockerfile/chart."),
        ("Assess security", "threat model → docs/security → Dockerfile → Helm security context/network policy → Kyverno → workflows."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Trace in this order"
    for task, trace in tasks:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = task, trace


def set_landscape_for_tables(doc: Document) -> None:
    """The document stays portrait; line tables are sized for portrait readability."""
    for section in doc.sections:
        if section.orientation == WD_ORIENT.LANDSCAPE:
            section.orientation = WD_ORIENT.PORTRAIT


def main() -> None:
    files = tracked_files()
    doc = Document()
    configure_document(doc)
    add_title(doc, files)
    add_toc(doc)
    add_scope(doc, files)
    add_system_map(doc)
    add_business_flow(doc)
    add_service_tour(doc)
    add_reading_order(doc)
    add_local_and_delivery(doc)
    add_inventory(doc, files)
    add_line_by_line(doc, files)
    add_appendix(doc)
    set_landscape_for_tables(doc)
    doc.core_properties.title = "Sales-App-SDLC Repository Reading Guide"
    doc.core_properties.subject = "Architecture, flow, reading order, and exhaustive line-by-line repository companion"
    doc.core_properties.author = "Sales-App-SDLC maintainers"
    doc.core_properties.keywords = "Sales-App-SDLC, architecture, microservices, line-by-line, GitOps, AKS"
    doc.core_properties.comments = "Generated by scripts/generate_repository_guide.py"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT.relative_to(ROOT)} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
